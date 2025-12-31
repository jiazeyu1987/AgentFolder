from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List

from core.util import ensure_dir, safe_read_text, sha256_file


def _extract_pdf(path: Path, *, max_chars: int) -> str:
    # Prefer PyMuPDF (fitz) for speed and robustness.
    try:
        import fitz  # type: ignore

        doc = fitz.open(str(path))
        parts: List[str] = []
        for page in doc:
            parts.append(page.get_text("text"))
            if sum(len(p) for p in parts) >= max_chars:
                break
        text = "\n".join(parts)
        return text[:max_chars]
    except Exception:
        pass

    # Fallback to pdfplumber if available.
    try:
        import pdfplumber  # type: ignore

        parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
                if sum(len(p) for p in parts) >= max_chars:
                    break
        text = "\n".join(parts)
        return text[:max_chars]
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("Missing dependency for PDF extraction (install PyMuPDF or pdfplumber)") from exc


def _extract_docx(path: Path, *, max_chars: int) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("Missing dependency for DOCX extraction (install python-docx)") from exc

    doc = Document(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
        if sum(len(x) for x in parts) >= max_chars:
            break

    # Also extract table text (common in regulatory docs).
    if sum(len(x) for x in parts) < max_chars:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text or "")
                    if sum(len(x) for x in parts) >= max_chars:
                        break
                if sum(len(x) for x in parts) >= max_chars:
                    break
            if sum(len(x) for x in parts) >= max_chars:
                break

    text = "\n".join(parts)
    return text[:max_chars]


def run(*, task_id: str, plan_id: str, inputs: List[Dict[str, Any]], params: Dict[str, Any]) -> Dict[str, Any]:
    max_chars = int(params.get("max_chars") or 200_000)
    artifacts: List[Dict[str, Any]] = []
    evidences: List[Dict[str, Any]] = []

    for inp in inputs:
        path = Path(str(inp.get("path") or ""))
        if not path.is_file():
            return {"status": "FAILED", "artifacts": [], "evidences": [], "error": {"code": "SKILL_BAD_INPUT", "message": f"not a file: {path}"}}

        suffix = path.suffix.lower()
        try:
            if suffix in {".txt", ".md"}:
                text = safe_read_text(path, max_chars=max_chars)
            elif suffix == ".pdf":
                text = _extract_pdf(path, max_chars=max_chars)
            elif suffix == ".docx":
                text = _extract_docx(path, max_chars=max_chars)
            else:
                return {
                    "status": "FAILED",
                    "artifacts": [],
                    "evidences": [],
                    "error": {"code": "SKILL_BAD_INPUT", "message": f"unsupported file type for text_extract: {suffix}"},
                }
        except Exception as exc:  # noqa: BLE001
            return {
                "status": "FAILED",
                "artifacts": [],
                "evidences": [],
                "error": {"code": "SKILL_FAILED", "message": f"{type(exc).__name__}: {exc}"},
            }

        sha = sha256_file(path)
        out_dir = Path("workspace") / "artifacts" / task_id
        ensure_dir(out_dir)
        out_path = out_dir / f"extracted_{sha}.txt"
        out_path.write_text(text, encoding="utf-8")
        artifacts.append({"name": "extracted_text", "path": str(out_path), "sha256": sha256_file(out_path), "format": "txt"})

    return {"status": "SUCCEEDED", "artifacts": artifacts, "evidences": evidences, "error": None}
